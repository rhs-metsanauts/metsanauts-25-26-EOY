import re
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.enum.style import WD_STYLE_TYPE

import sys
MD_FILE   = sys.argv[1] if len(sys.argv) > 1 else r"c:\Users\busyp\final\METSAnauts_Technical_Reference.md"
DOCX_FILE = sys.argv[2] if len(sys.argv) > 2 else r"c:\Users\busyp\final\METSAnauts_Technical_Reference.docx"

# ── helpers ──────────────────────────────────────────────────────────────────

def add_hyperlink(paragraph, text, url):
    part = paragraph.part
    r_id = part.relate_to(url, "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink", is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    new_run = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rStyle = OxmlElement("w:rStyle")
    rStyle.set(qn("w:val"), "Hyperlink")
    rPr.append(rStyle)
    new_run.append(rPr)
    new_run.text = text
    hyperlink.append(new_run)
    paragraph._p.append(hyperlink)


def set_cell_bg(cell, hex_color):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


INLINE_RE = re.compile(r"(\*\*(.+?)\*\*|`(.+?)`|\[([^\]]+)\]\(([^)]+)\))")

def add_inline(paragraph, text):
    """Parse bold, inline code, and links within a text segment."""
    pos = 0
    for m in INLINE_RE.finditer(text):
        # plain text before match
        if m.start() > pos:
            paragraph.add_run(text[pos:m.start()])
        full = m.group(0)
        if full.startswith("**"):
            run = paragraph.add_run(m.group(2))
            run.bold = True
        elif full.startswith("`"):
            run = paragraph.add_run(m.group(3))
            run.font.name = "Courier New"
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0xC7, 0x25, 0x4E)
        else:
            add_hyperlink(paragraph, m.group(4), m.group(5))
        pos = m.end()
    if pos < len(text):
        paragraph.add_run(text[pos:])


# ── document setup ────────────────────────────────────────────────────────────

doc = Document()

# Page margins
for section in doc.sections:
    section.top_margin    = Cm(2.0)
    section.bottom_margin = Cm(2.0)
    section.left_margin   = Cm(2.5)
    section.right_margin  = Cm(2.5)

# Normal style baseline
normal = doc.styles["Normal"]
normal.font.name = "Calibri"
normal.font.size = Pt(10.5)

# Heading styles
HEADING_SIZES = {1: 20, 2: 15, 3: 12, 4: 11}
HEADING_COLORS = {
    1: RGBColor(0x1F, 0x4E, 0x79),
    2: RGBColor(0x27, 0x72, 0xB4),
    3: RGBColor(0x2E, 0x86, 0xAB),
    4: RGBColor(0x1A, 0x1A, 0x2E),
}
for lvl in range(1, 5):
    style = doc.styles[f"Heading {lvl}"]
    style.font.name = "Calibri"
    style.font.size = Pt(HEADING_SIZES[lvl])
    style.font.color.rgb = HEADING_COLORS[lvl]
    style.font.bold = True
    pf = style.paragraph_format
    pf.space_before = Pt(12 if lvl <= 2 else 8)
    pf.space_after  = Pt(4)

# Ensure Hyperlink character style exists
if "Hyperlink" not in [s.name for s in doc.styles]:
    hl = doc.styles.add_style("Hyperlink", WD_STYLE_TYPE.CHARACTER)
    hl.font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    hl.font.underline = True
else:
    doc.styles["Hyperlink"].font.color.rgb = RGBColor(0x05, 0x63, 0xC1)
    doc.styles["Hyperlink"].font.underline = True


# ── parsing ───────────────────────────────────────────────────────────────────

with open(MD_FILE, encoding="utf-8") as f:
    lines = f.readlines()

i = 0
while i < len(lines):
    raw = lines[i].rstrip("\n")
    stripped = raw.strip()

    # skip bare horizontal rules
    if re.match(r"^-{3,}$", stripped):
        i += 1
        continue

    # headings
    hm = re.match(r"^(#{1,4})\s+(.*)", stripped)
    if hm:
        level = len(hm.group(1))
        text  = hm.group(2).strip()
        # strip markdown links from heading text for simplicity
        text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)
        text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
        doc.add_heading(text, level=level)
        i += 1
        continue

    # fenced code block
    if stripped.startswith("```"):
        i += 1
        code_lines = []
        while i < len(lines) and not lines[i].strip().startswith("```"):
            code_lines.append(lines[i].rstrip("\n"))
            i += 1
        i += 1  # consume closing ```
        p = doc.add_paragraph()
        p.style = "Normal"
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.space_after  = Pt(4)
        run = p.add_run("\n".join(code_lines))
        run.font.name = "Courier New"
        run.font.size = Pt(8.5)
        # light grey shading on paragraph
        pPr = p._p.get_or_add_pPr()
        shd = OxmlElement("w:shd")
        shd.set(qn("w:val"),   "clear")
        shd.set(qn("w:color"), "auto")
        shd.set(qn("w:fill"),  "F2F2F2")
        pPr.append(shd)
        continue

    # blockquote
    if stripped.startswith(">"):
        text = stripped.lstrip("> ").strip()
        p = doc.add_paragraph()
        p.style = "Normal"
        p.paragraph_format.left_indent  = Inches(0.4)
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(2)
        run = p.add_run(text)
        run.italic = True
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
        i += 1
        continue

    # table — collect all rows
    if stripped.startswith("|"):
        table_lines = []
        while i < len(lines) and lines[i].strip().startswith("|"):
            table_lines.append(lines[i].strip())
            i += 1

        # filter separator rows
        rows = [r for r in table_lines if not re.match(r"^\|[-| :]+\|$", r)]
        if not rows:
            continue

        def split_row(r):
            return [c.strip() for c in r.strip("|").split("|")]

        cols = len(split_row(rows[0]))
        tbl  = doc.add_table(rows=len(rows), cols=cols)
        tbl.style = "Table Grid"

        for ri, row_text in enumerate(rows):
            cells = split_row(row_text)
            for ci, cell_text in enumerate(cells[:cols]):
                cell = tbl.cell(ri, ci)
                cell.text = ""
                p = cell.paragraphs[0]
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after  = Pt(2)
                add_inline(p, cell_text)
                if ri == 0:
                    for run in p.runs:
                        run.bold = True
                    set_cell_bg(cell, "1F4E79")
                    for run in p.runs:
                        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                elif ri % 2 == 0:
                    set_cell_bg(cell, "DCE6F1")
        doc.add_paragraph()
        continue

    # bullet list
    if stripped.startswith("- "):
        text = stripped[2:].strip()
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.space_after  = Pt(1)
        add_inline(p, text)
        i += 1
        continue

    # blank line
    if not stripped:
        i += 1
        continue

    # normal paragraph
    p = doc.add_paragraph()
    p.style = "Normal"
    p.paragraph_format.space_after = Pt(6)
    add_inline(p, stripped)
    i += 1

doc.save(DOCX_FILE)
print(f"Saved: {DOCX_FILE}")
